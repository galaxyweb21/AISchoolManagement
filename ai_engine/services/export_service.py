# ai_engine/services/export_service.py
"""
Export Service for AI-generated content
Supports PDF and DOC document generation
"""

import io
import os
from datetime import datetime
from django.conf import settings
from django.http import HttpResponse
from django.template.loader import render_to_string
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import tempfile
import logging

logger = logging.getLogger(__name__)

# Try to import xhtml2pdf, fallback to a warning if not available
try:
    from xhtml2pdf import pisa

    XHTML2PDF_AVAILABLE = True
except ImportError as e:
    XHTML2PDF_AVAILABLE = False
    logger.warning(f"xhtml2pdf not available: {e}. PDF export will use fallback.")


class ExportService:
    """Service for exporting AI-generated content to PDF and DOC formats"""

    @staticmethod
    def _pdf_link_callback(uri, rel):
        from urllib.parse import urlparse
        path = urlparse(uri).path if uri else ''
        media_url = getattr(settings, 'MEDIA_URL', '/media/')
        static_url = getattr(settings, 'STATIC_URL', '/static/')
        if path.startswith(media_url):
            candidate = os.path.join(getattr(settings, 'MEDIA_ROOT', ''), path[len(media_url):].lstrip('/'))
        elif path.startswith(static_url):
            candidate = os.path.join(getattr(settings, 'STATIC_ROOT', ''), path[len(static_url):].lstrip('/'))
            if not os.path.exists(candidate):
                candidate = os.path.join(getattr(settings, 'BASE_DIR', ''), 'static', path[len(static_url):].lstrip('/'))
        else:
            candidate = uri
        return candidate if candidate and os.path.exists(candidate) else uri

    @staticmethod
    def _generate_pdf(html_content, filename):
        if not XHTML2PDF_AVAILABLE:
            return ExportService._fallback_pdf_response(filename)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        status = pisa.CreatePDF(html_content, dest=response, encoding='UTF-8', link_callback=ExportService._pdf_link_callback)
        if status.err:
            logger.error('xhtml2pdf returned %s error(s) for %s', status.err, filename)
            return ExportService._fallback_pdf_response(filename)
        return response

    @staticmethod
    def _fallback_pdf_response(filename):
        """
        Fallback when xhtml2pdf is not available.
        Returns a simple HTML page with instructions.
        """
        fallback_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>PDF Export Unavailable</title>
            <style>
                body {{ font-family: Arial, sans-serif; text-align: center; padding: 50px; }}
                .container {{ max-width: 600px; margin: 0 auto; }}
                .icon {{ font-size: 64px; color: #dc3545; }}
                h1 {{ color: #333; }}
                p {{ color: #666; line-height: 1.6; }}
                .btn {{ display: inline-block; padding: 12px 24px; background: #0d6efd; color: white; text-decoration: none; border-radius: 6px; }}
                .btn:hover {{ background: #0a58ca; }}
                .options {{ text-align: left; max-width: 400px; margin: 20px auto; }}
                .options li {{ margin-bottom: 8px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="icon">📄</div>
                <h1>PDF Export Unavailable</h1>
                <p>The PDF export feature requires the xhtml2pdf library which is not installed.</p>
                <p><strong>Alternative Options:</strong></p>
                <ul class="options">
                    <li>✅ Use the <strong>DOC</strong> export option (Microsoft Word format)</li>
                    <li>✅ Use the <strong>Print</strong> function in your browser and select "Save as PDF"</li>
                    <li>📦 Install xhtml2pdf: <code>pip install xhtml2pdf</code></li>
                </ul>
                <p style="margin-top: 30px;">
                    <a href="#" onclick="window.print(); return false;" class="btn">🖨️ Print this page (Save as PDF)</a>
                </p>
                <p style="margin-top: 20px; font-size: 12px; color: #999;">
                    Content for: {filename}
                </p>
            </div>
        </body>
        </html>
        """
        response = HttpResponse(fallback_html, content_type='text/html')
        response['Content-Disposition'] = f'attachment; filename="{filename.replace(".pdf", ".html")}"'
        return response

    @staticmethod
    def _get_pdf_styles():
        """Get common PDF styles"""
        return """
        <style>
            body {
                font-family: 'Helvetica', 'Arial', sans-serif;
                font-size: 12px;
                line-height: 1.6;
                margin: 40px;
                color: #333;
            }
            .header {
                text-align: center;
                margin-bottom: 30px;
                border-bottom: 2px solid #0d6efd;
                padding-bottom: 20px;
            }
            .header h1 {
                font-size: 24px;
                color: #0d6efd;
                margin: 0;
            }
            .header h2 {
                font-size: 18px;
                color: #333;
                margin: 5px 0;
            }
            .header .subtitle {
                color: #666;
                font-size: 12px;
            }
            .meta {
                background: #f8f9fa;
                padding: 15px;
                border-radius: 5px;
                margin-bottom: 20px;
            }
            .meta table {
                width: 100%;
            }
            .meta td {
                padding: 3px 10px;
            }
            .meta td:first-child {
                font-weight: bold;
                width: 120px;
            }
            .question {
                margin-bottom: 20px;
                padding-bottom: 15px;
                border-bottom: 1px solid #eee;
            }
            .question .q-header {
                font-weight: bold;
                font-size: 13px;
                color: #0d6efd;
            }
            .question .q-text {
                margin: 8px 0;
            }
            .question .q-points {
                color: #666;
                font-size: 11px;
            }
            .question .options {
                margin: 5px 0 5px 20px;
            }
            .question .options li {
                list-style-type: disc;
                margin-bottom: 2px;
            }
            .footer {
                text-align: center;
                margin-top: 40px;
                padding-top: 20px;
                border-top: 1px solid #ddd;
                color: #999;
                font-size: 10px;
            }
            .instructions {
                background: #f0f7ff;
                padding: 12px 15px;
                border-radius: 5px;
                margin-bottom: 20px;
                border-left: 4px solid #0d6efd;
            }
            .answer-space {
                min-height: 60px;
                border-bottom: 1px dashed #ccc;
                margin: 5px 0 10px 0;
            }
            .section-title {
                font-size: 16px;
                font-weight: bold;
                color: #198754;
                margin: 20px 0 10px 0;
                border-bottom: 1px solid #198754;
                padding-bottom: 5px;
            }
            .subject-table {
                width: 100%;
                border-collapse: collapse;
                margin: 10px 0;
            }
            .subject-table th {
                background: #198754;
                color: white;
                padding: 8px 12px;
                text-align: left;
            }
            .subject-table td {
                padding: 6px 12px;
                border-bottom: 1px solid #ddd;
            }
            .subject-table tr:nth-child(even) {
                background: #f8f9fa;
            }
            .comment-box {
                background: #f8f9fa;
                padding: 15px;
                border-radius: 5px;
                margin: 15px 0;
                border-left: 4px solid #198754;
            }
            .overall-score {
                font-size: 18px;
                font-weight: bold;
                color: #198754;
            }
            .risk-critical { color: #dc3545; font-weight: bold; }
            .risk-high { color: #ffc107; font-weight: bold; }
            .risk-medium { color: #0dcaf0; font-weight: bold; }
            .risk-low { color: #198754; font-weight: bold; }
            .risk-factors li {
                margin-bottom: 5px;
            }
            .factor-points {
                float: right;
                font-weight: bold;
                color: #6c757d;
            }
        </style>
        """

    @staticmethod
    def export_exam_to_pdf(exam, questions):
        """Export an AI-generated exam to PDF"""
        context = {
            'exam': exam,
            'questions': questions,
            'total_points': sum(q.points for q in questions),
            'generated_at': datetime.now(),
            'school_name': exam.school.name if exam.school else 'School',
        }

        html_content = render_to_string('ai_engine/exports/exam_export_pdf.html', context)

        # Wrap with styles
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{exam.title}</title>
            {ExportService._get_pdf_styles()}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        filename = f"exam_{exam.title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        return ExportService._generate_pdf(full_html, filename)

    @staticmethod
    def export_exam_to_doc(exam, questions):
        """Export an AI-generated exam to DOC (Word) format"""
        doc = Document()

        # Add school header
        header = doc.add_heading(exam.school.name if exam.school else 'School', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Exam title
        title = doc.add_heading(exam.title, 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Exam metadata
        doc.add_paragraph(f"Subject: {exam.subject}")
        doc.add_paragraph(f"Class: {exam.school_class.name if exam.school_class else 'N/A'}")
        doc.add_paragraph(f"Grade Level: {exam.grade_level}")
        doc.add_paragraph(f"Difficulty: {exam.get_difficulty_display()}")
        doc.add_paragraph(f"Total Points: {sum(q.points for q in questions)}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        doc.add_paragraph("Instructions: Answer all questions. Read each question carefully.")
        doc.add_paragraph("")

        # Add questions
        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph()
            p.add_run(f"Q{i}. ").bold = True
            p.add_run(q.question_text)
            doc.add_paragraph(f"[{q.points} point{'s' if q.points > 1 else ''}]")

            if q.question_type == 'MCQ' and q.options:
                for opt in q.options:
                    doc.add_paragraph(f"   • {opt}", style='List Bullet')

            if q.question_type in ['SHORT_ANSWER', 'ESSAY']:
                for _ in range(3 if q.question_type == 'SHORT_ANSWER' else 6):
                    doc.add_paragraph("")

            doc.add_paragraph("")

        doc.add_paragraph("")
        doc.add_paragraph("--- End of Exam ---")
        doc.add_paragraph(f"Generated by EduAI on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"exam_{exam.title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    @staticmethod
    def _report_card_data(report_card):
        '''Use live TerminalResult data for draft cards. Finalized cards keep their
        locked official average/grade/position, but any subject row still missing
        Class /30 or Exam /70 is repaired from the live TerminalResult data so an
        export never shows a dash once the marks have actually been entered.'''
        from ai_engine.services.report_card_engine import ReportCardEngine
        report_card, computed = ReportCardEngine.refresh_report_card_snapshot(report_card, save=True)
        return computed

    @staticmethod
    def _safe_file_url(field):
        try:
            return field.url if field else ''
        except Exception:
            return ''

    @staticmethod
    def export_report_card_to_pdf(report_card):
        school, student = report_card.school, report_card.student
        computed = ExportService._report_card_data(report_card)
        context = {
            'report_card': report_card,
            'student': student,
            'academic_term': report_card.academic_term,
            'subject_breakdown': computed.get('subject_breakdown', []),
            'generated_at': datetime.now(),
            'school_logo': ExportService._safe_file_url(getattr(school, 'logo', None)),
            'student_photo': ExportService._safe_file_url(getattr(student, 'profile_photo', None)),
        }
        html = render_to_string('ai_engine/exports/report_card_export_pdf.html', context)
        css = '''<style>
        @page{size:A4;margin:8mm 9mm 10mm 9mm}
        body{font-family:Helvetica,Arial,sans-serif;color:#20312d;font-size:8.6pt;margin:0;background:#fff}
        table{width:100%;border-collapse:collapse}
        .sheet{border:1.5px solid #e2e6e2;border-radius:4px;padding:2px}

        /* header band */
        .band{background-color:#17352c;padding:0}
        .band-logo{width:64px;padding:10px 4px 10px 12px;vertical-align:middle}
        .logo{width:50px;height:50px}
        .band-mid{padding:9px 10px;vertical-align:middle}
        .school-name{color:#ffffff;font-size:15pt;font-weight:bold;letter-spacing:.3px}
        .school-sub{color:#bcd0c8;font-size:7.3pt;margin-top:1px}
        .title-pill{display:inline;color:#17352c;background-color:#c9962b;font-size:9pt;font-weight:bold;padding:3px 10px;border-radius:9px;margin-top:6px}
        .subtitle{color:#e7c88a;font-size:7.3pt;font-weight:bold;letter-spacing:.4px;margin-top:5px}
        .band-photo{width:92px;padding:8px 12px 8px 4px;text-align:right;vertical-align:middle}
        .photo-frame{background-color:#c9962b;border-radius:8px;display:inline-block}
        .photo{width:72px;height:84px;border-radius:6px}
        .photo-fallback{width:72px;height:84px;line-height:84px;text-align:center;color:#ffffff;font-size:18pt;font-weight:bold;background-color:#2c7a6d;border-radius:6px}

        /* info + kpis */
        .info td{border:1px solid #e2e6e2;padding:5px 7px}
        .info .label{background-color:#f5f2e8;font-weight:bold;color:#5c6a63;width:14%}
        .info .val{width:36%}
        .kpis td{padding:0 3px}
        .kpi{border-radius:8px;text-align:center;padding:7px 2px}
        .kpi-label{font-size:6.6pt;font-weight:bold;letter-spacing:.5px;color:#5c6a63;text-transform:uppercase}
        .kpi-value{font-size:12pt;font-weight:bold;margin-top:2px}
        .kpi-value-sm{font-size:8.5pt}
        .kpi-gold{background-color:#f6ecd9}
        .kpi-gold .kpi-value{color:#946b1c}
        .kpi-green{background-color:#e4efe9}
        .kpi-green .kpi-value{color:#17352c}
        .kpi-teal{background-color:#dcefec}
        .kpi-teal .kpi-value{color:#2c7a6d}

        /* section labels */
        .section{background-color:#17352c;color:#ffffff;font-weight:bold;font-size:8.6pt;letter-spacing:.4px;padding:5px 8px;border-radius:5px;margin-top:10px;margin-bottom:5px}
        .section-dot{display:inline-block;width:6px;height:6px;background-color:#c9962b;border-radius:3px;margin-right:6px}

        /* results table */
        .results th{background-color:#946b1c;color:#ffffff;padding:6px 5px;font-size:7.6pt;text-transform:uppercase;letter-spacing:.3px}
        .results td{border-bottom:1px solid #e9ece9;padding:5.5px 5px}
        .col-subject{text-align:left;font-weight:bold;width:22%}
        .col-remark{text-align:left;color:#5c6a63}
        .center{text-align:center}
        .r-a{background-color:#ffffff}
        .r-b{background-color:#f8f9f6}
        .chip{display:inline-block;min-width:38px;padding:2.5px 7px;border-radius:8px;font-weight:bold}
        .chip-teal{background-color:#dcefec;color:#1f5c53}
        .chip-gold{background-color:#f6ecd9;color:#7a5917}
        .grade-cell{font-weight:bold;color:#17352c}
        .dash{color:#a3aca7}
        .scale-note{color:#7a857f;font-size:7pt;font-style:italic;margin-top:4px}

        /* comments */
        .comments td{width:50%;padding:0 4px 0 0;vertical-align:top}
        .comment-title{font-weight:bold;color:#17352c;font-size:8pt;margin-bottom:3px}
        .comment-box{border:1px solid #e2e6e2;border-radius:6px;background-color:#fbfcfb;padding:8px;min-height:34px;line-height:1.4}

        .sign td{text-align:center;height:40px;vertical-align:bottom;font-size:8pt}
        .footer{text-align:center;color:#8a938d;font-size:6.6pt;border-top:1px solid #e2e6e2;margin-top:9px;padding-top:4px}
        </style>'''
        full = '<!DOCTYPE html><html><head><meta charset="utf-8"><title>End-of-Term Report Card</title>'+css+'</head><body>'+html+'</body></html>'
        return ExportService._generate_pdf(full, f'report_card_{student.admission_number}_{datetime.now():%Y%m%d}.pdf')

    @staticmethod
    def _shade_cell(cell, fill):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd'); tcPr.append(shd)
        shd.set(qn('w:fill'), fill)

    @staticmethod
    def _shade_paragraph_bg(paragraph, fill):
        """Shade a plain paragraph's background (used for section-header bars
        that aren't inside a table cell)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        pPr.append(shd)

    @staticmethod
    def _set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
        for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            node = tcMar.find(qn(f'w:{m}'))
            if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
            node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

    @staticmethod
    def export_report_card_to_doc(report_card):
        school, student = report_card.school, report_card.student
        computed = ExportService._report_card_data(report_card)
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin=Inches(.42); sec.bottom_margin=Inches(.45); sec.left_margin=Inches(.5); sec.right_margin=Inches(.5)
        styles=doc.styles
        styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9)

        head=doc.add_table(rows=1, cols=3); head.autofit=False
        col_widths=[Inches(.75), Inches(5.35), Inches(1.4)]
        for i,w in enumerate(col_widths):
            head.columns[i].width=w
            for row in head.rows: row.cells[i].width=w
        for c in head.rows[0].cells:
            ExportService._shade_cell(c,'17352C'); c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        if getattr(school,'logo',None) and getattr(school.logo,'path',None) and os.path.exists(school.logo.path):
            head.cell(0,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            head.cell(0,0).paragraphs[0].add_run().add_picture(school.logo.path,width=Inches(.6))
        if getattr(student,'profile_photo',None) and getattr(student.profile_photo,'path',None) and os.path.exists(student.profile_photo.path):
            photo_cell=head.cell(0,2); photo_cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            photo_cell.paragraphs[0].add_run().add_picture(student.profile_photo.path,width=Inches(1.0),height=Inches(1.15))
        else:
            photo_cell=head.cell(0,2); p0=photo_cell.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
            initials=(student.user.first_name[:1]+student.user.last_name[:1]).upper() if student.user.first_name and student.user.last_name else '—'
            rr=p0.add_run(initials); rr.bold=True; rr.font.size=Pt(18); rr.font.color.rgb=RGBColor(255,255,255)
        p=head.cell(0,1).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        r=p.add_run(school.name.upper()); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(255,255,255)
        contact_bits=[str(b) for b in [getattr(school,'address',''), getattr(school,'phone_number',''), getattr(school,'contact_email','')] if b]
        if contact_bits:
            r2=p.add_run('\n'+' \u2022 '.join(contact_bits)); r2.font.size=Pt(7.5); r2.font.color.rgb=RGBColor(188,208,200)
        for c in head.rows[0].cells: ExportService._set_cell_margins(c,80,60,80,60)

        # gold title pill, centered, under the header banner
        pill=doc.add_table(rows=1,cols=1); pill.autofit=False; pill.columns[0].width=Inches(3.2)
        pill.alignment=WD_TABLE_ALIGNMENT.CENTER
        pill_cell=pill.cell(0,0); ExportService._shade_cell(pill_cell,'C9962B'); ExportService._set_cell_margins(pill_cell,45,80,45,80)
        pp=pill_cell.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pr=pp.add_run('END-OF-TERM REPORT CARD'); pr.bold=True; pr.font.size=Pt(11); pr.font.color.rgb=RGBColor(23,53,44)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=p.add_run(f'{report_card.academic_term}  •  {"FINALIZED & OFFICIAL" if report_card.is_finalized else "DRAFT"}'); rr.bold=True; rr.font.size=Pt(8.5); rr.font.color.rgb=RGBColor(90,100,95)

        info=doc.add_table(rows=2,cols=4); info.style='Table Grid'
        vals=[('Student',student.user.get_full_name(),'Admission No.',student.admission_number),('Grade',str(student.grade_level or 'N/A'),'Class',str(student.school_class or 'N/A'))]
        for i,row in enumerate(vals):
            for j,v in enumerate(row):
                cell=info.cell(i,j); cell.text=str(v); ExportService._set_cell_margins(cell)
                if j in (0,2): ExportService._shade_cell(cell,'F3F1E9'); cell.paragraphs[0].runs[0].bold=True

        summary=doc.add_table(rows=2,cols=5); summary.style='Table Grid'
        avg=report_card.overall_average
        labels=['AVERAGE','GRADE','POSITION','ATTENDANCE','PROMOTION']
        vals=[f'{avg:.1f}%' if avg is not None else '—', report_card.overall_grade or '—', f'{report_card.overall_position}/{report_card.class_size}' if report_card.overall_position else '—', f'{report_card.attendance_rate:.1f}%' if report_card.attendance_rate is not None else '—', report_card.promotion_status or '—']
        kpi_colors=['F6ECD9','E4EFE9','DCEFEC','F6ECD9','E4EFE9']
        kpi_text_colors=[RGBColor(148,107,28),RGBColor(23,53,44),RGBColor(44,122,109),RGBColor(148,107,28),RGBColor(23,53,44)]
        for j,(label,v,fill,tc) in enumerate(zip(labels,vals,kpi_colors,kpi_text_colors)):
            lc=summary.cell(0,j); ExportService._shade_cell(lc,fill); lc.text=label; lc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            lc.paragraphs[0].runs[0].bold=True; lc.paragraphs[0].runs[0].font.size=Pt(6.8); lc.paragraphs[0].runs[0].font.color.rgb=RGBColor(90,100,95)
            vc=summary.cell(1,j); ExportService._shade_cell(vc,fill); vc.text=str(v); vc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            vc.paragraphs[0].runs[0].bold=True; vc.paragraphs[0].runs[0].font.size=Pt(11); vc.paragraphs[0].runs[0].font.color.rgb=tc
            ExportService._set_cell_margins(lc,50,40,20,40); ExportService._set_cell_margins(vc,20,40,60,40)

        p=doc.add_paragraph(); r=p.add_run('  ACADEMIC RESULTS'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(255,255,255)
        p.paragraph_format.space_before=Pt(10)
        ExportService._shade_paragraph_bg(p, '17352C')
        table=doc.add_table(rows=1, cols=6); table.style='Table Grid'
        headers=['Subject','Class /30','Exam /70','Final /100','Grade','Remark']
        for c,t in zip(table.rows[0].cells,headers):
            c.text=t; ExportService._shade_cell(c,'946B1C'); c.paragraphs[0].runs[0].bold=True; c.paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for row_idx,item in enumerate(computed.get('subject_breakdown',[])):
            cells=table.add_row().cells
            cs=item.get('class_score'); es=item.get('exam_score'); total=item.get('total')
            vals=[item.get('subject',''),f'{float(cs):.2f}' if cs is not None else '—',f'{float(es):.2f}' if es is not None else '—',f'{float(total):.2f}' if total is not None else '—',item.get('grade','—') or '—',item.get('remark','') or '—']
            zebra = row_idx % 2 == 1
            for idx,(cell,v) in enumerate(zip(cells,vals)):
                cell.text=str(v); ExportService._set_cell_margins(cell,50,60,50,60); cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if idx in (0,5) else WD_ALIGN_PARAGRAPH.CENTER
                if zebra: ExportService._shade_cell(cell,'F8F9F6')
                if idx in (1,2): cell.paragraphs[0].runs[0].bold=True; cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(31,92,83)
                if idx==3: cell.paragraphs[0].runs[0].bold=True; cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(148,107,28)
                if idx==4: cell.paragraphs[0].runs[0].bold=True; cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(23,53,44)
        p=doc.add_paragraph('Official terminal scale: Class /30 + Examination /70 = Final /100.'); p.runs[0].italic=True; p.runs[0].font.size=Pt(8)

        p=doc.add_paragraph(); r=p.add_run('  ATTENDANCE & GENERAL ASSESSMENT'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(255,255,255)
        p.paragraph_format.space_before=Pt(6)
        ExportService._shade_paragraph_bg(p, '17352C')
        att=doc.add_table(rows=2,cols=6); att.style='Table Grid'
        av=[('Present',report_card.attendance_present),('Late',report_card.attendance_late),('Absent',report_card.attendance_absent),('Recorded',report_card.attendance_total),('Conduct',report_card.conduct or '—'),('Promotion',report_card.promotion_status or '—')]
        for i,(label,value) in enumerate(av):
            att.cell(0,i).text=label; att.cell(1,i).text=str(value); ExportService._shade_cell(att.cell(0,i),'F3F1E9'); att.cell(0,i).paragraphs[0].runs[0].bold=True; att.cell(0,i).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; att.cell(1,i).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

        for title,text in [("CLASS TEACHER'S COMMENT", report_card.teacher_comment or report_card.ai_narrative or 'No comment recorded.'),("HEADTEACHER'S COMMENT",report_card.headteacher_comment or 'No comment recorded.')]:
            p=doc.add_paragraph(); r=p.add_run(title); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(23,53,44)
            box=doc.add_table(rows=1,cols=1); box.style='Table Grid'; box.cell(0,0).text=text; ExportService._shade_cell(box.cell(0,0),'FBFCFB'); ExportService._set_cell_margins(box.cell(0,0),100,100,100,100)

        sig=doc.add_table(rows=2,cols=3); sig.style='Table Grid'
        for i,t in enumerate(['Class Teacher','Headteacher','Parent / Guardian']):
            sig.cell(0,i).text='\n\n________________________'; sig.cell(1,i).text=t; sig.cell(1,i).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; sig.cell(0,i).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        foot=sec.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.CENTER; foot.text=f'{school.name} • {report_card.academic_term} • EduAI School Management'
        for run in foot.runs: run.font.size=Pt(7); run.font.color.rgb=RGBColor(110,120,115)
        response=HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition']=f'attachment; filename="report_card_{student.admission_number}_{datetime.now():%Y%m%d}.docx"'
        doc.save(response); return response

    @staticmethod
    def export_risk_assessment_to_pdf(assessment):
        """Export a risk assessment to PDF"""
        context = {
            'assessment': assessment,
            'student': assessment.student,
            'generated_at': datetime.now(),
        }

        html_content = render_to_string('ai_engine/exports/risk_assessment_export_pdf.html', context)

        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Risk Assessment - {assessment.student.user.get_full_name()}</title>
            {ExportService._get_pdf_styles()}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        filename = f"risk_assessment_{assessment.student.user.get_full_name()}_{datetime.now().strftime('%Y%m%d')}.pdf"
        return ExportService._generate_pdf(full_html, filename)

    @staticmethod
    def export_risk_assessment_to_doc(assessment):
        """Export a risk assessment to DOC (Word) format"""
        doc = Document()

        header = doc.add_heading(assessment.school.name if assessment.school else 'School', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title = doc.add_heading('Student Risk Assessment Report', 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Student: {assessment.student.user.get_full_name()}")
        doc.add_paragraph(f"Admission Number: {assessment.student.admission_number}")
        doc.add_paragraph(f"Grade Level: {assessment.student.grade_level}")
        doc.add_paragraph(
            f"Class: {assessment.student.school_class.name if assessment.student.school_class else 'N/A'}")
        doc.add_paragraph(f"Assessment Date: {assessment.run.computed_at if assessment.run else datetime.now()}")
        doc.add_paragraph("")

        doc.add_heading('Risk Assessment Summary', 2)
        doc.add_paragraph(f"Risk Score: {assessment.risk_score}/100")
        doc.add_paragraph(f"Risk Band: {assessment.risk_band}")
        doc.add_paragraph("")

        doc.add_heading('Contributing Factors', 2)
        if assessment.contributing_factors:
            for factor in assessment.contributing_factors:
                doc.add_paragraph(f"• {factor['factor']}: {factor['detail']} (Points: {factor['points']})")
        else:
            doc.add_paragraph("No specific risk factors identified.")

        doc.add_paragraph("")

        doc.add_heading('Detailed Analysis', 2)
        doc.add_paragraph(
            f"Attendance Rate: {assessment.attendance_rate}%" if assessment.attendance_rate else "Attendance: N/A")
        doc.add_paragraph(
            f"Grade Average: {assessment.grade_average}%" if assessment.grade_average else "Grade Average: N/A")
        doc.add_paragraph(
            f"Fee Overdue: {assessment.fee_overdue_amount}" if assessment.fee_overdue_amount else "Fee Overdue: None")
        doc.add_paragraph("")

        if assessment.narrative:
            doc.add_heading('Narrative Summary', 2)
            doc.add_paragraph(assessment.narrative)
            doc.add_paragraph("")

        doc.add_paragraph("--- End of Risk Assessment ---")
        doc.add_paragraph(f"Generated by EduAI on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"risk_assessment_{assessment.student.user.get_full_name()}_{datetime.now().strftime('%Y%m%d')}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    @staticmethod
    def export_finance_insight_to_pdf(snapshot, assessments):
        """Export finance insights to PDF"""
        context = {
            'snapshot': snapshot,
            'assessments': assessments,
            'generated_at': datetime.now(),
            'school_name': snapshot.get('school_name', 'School'),
        }

        html_content = render_to_string('ai_engine/exports/finance_insight_export_pdf.html', context)

        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Finance Insights</title>
            {ExportService._get_pdf_styles()}
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        filename = f"finance_insights_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        return ExportService._generate_pdf(full_html, filename)